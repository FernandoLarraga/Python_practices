# Im trying to make automated task
# Create a document every day with the date

import shutil
from datetime import datetime
from importlib import invalidate_caches
now = datetime.now() 
from docx import Document
from  docx.shared import Inches

year = str(now.year)
month = str(now.month) 
day = str(now.day)
hour = str(now.hour)
minutes = str(now.minute)

date = year + "-" +  month + "-" + day 
hour_minutes = hour + ":" + minutes

document = Document()

document.add_heading('Automated document')

p = document.add_paragraph('This is a automated document where I`m going to put the')
p.add_run(' date').bold = True

p = document.add_paragraph("This is the date: ")
p.add_run(date)
p = document.add_paragraph("This is the hour: ")
p.add_run(hour_minutes)

document_name = "Date_time_" + date + ".docx"
document.add_page_break()
document.save(document_name)

original = r"C:\Users\ferna\OneDrive\Documents\Visual Studio 2022\Python\{document_name}".format(document_name = document_name)
target = r"C:\Users\ferna\Downloads\Pruebas_python\{document_name}".format(document_name = document_name) 

#C:\Users\ferna\OneDrive\Documents\Visual Studio 2022\Python

shutil.move(original, target)


